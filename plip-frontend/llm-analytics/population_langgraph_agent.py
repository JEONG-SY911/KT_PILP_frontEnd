####### 인구 데이터 분석용 LangGraph Agent with RAG & Document Support #######
## 1. 라이브러리 로딩 ---------------------------------------------
import pandas as pd
import numpy as np
import os
import openai
import random
import ast
from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime
from typing import List, Dict, Any, Literal, TypedDict
import json
import logging
import warnings
from dotenv import load_dotenv

# .env 파일 로드
load_dotenv()

warnings.filterwarnings("ignore", category=DeprecationWarning)

from typing import Annotated, Literal, Sequence, TypedDict, List, Dict
from langchain import hub
from langchain_core.messages import BaseMessage, HumanMessage
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import PromptTemplate, ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.output_parsers import CommaSeparatedListOutputParser
from langgraph.graph import StateGraph, END

# RAG 관련 라이브러리 추가
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader, DirectoryLoader
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor
import chromadb
from chromadb.config import Settings

# Word 문서 처리 라이브러리 추가
from docx import Document
import fitz  # PyMuPDF
import tempfile
import shutil

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 환경변수에서 설정 로드
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
HOST = os.getenv("HOST", "0.0.0.0")
PORT = 8004  # Population LangGraph 전용 포트
ENVIRONMENT = os.getenv("ENVIRONMENT", "development")

# API 키 확인
if not OPENAI_API_KEY:
    logger.error("OPENAI_API_KEY가 설정되지 않았습니다. .env 파일을 확인해주세요.")
    raise ValueError("OPENAI_API_KEY가 필요합니다.")

logger.info(f"서버 환경: {ENVIRONMENT}")
logger.info(f"OpenAI API 키 설정됨: {'*' * 10}{OPENAI_API_KEY[-4:] if OPENAI_API_KEY else 'None'}")

app = FastAPI(title="인구 데이터 LangGraph Agent 분석 API with RAG", version="2.0.0")

# CORS 설정
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

## ---------------- RAG 시스템 초기화 ----------------------

class RAGSystem:
    def __init__(self):
        self.vectorstore = None
        self.embeddings = None
        self.retriever = None
        self.initialize_rag()
    
    def initialize_rag(self):
        """RAG 시스템 초기화"""
        try:
            # 한국어 지원 임베딩 모델 사용
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            
            # 벡터 스토어 초기화
            self.vectorstore = Chroma(
                collection_name="population_knowledge",
                embedding_function=self.embeddings,
                persist_directory="./chroma_db"
            )
            
            # 기본 지식 베이스 구축
            self.build_knowledge_base()
            
            # 검색기 설정
            self.retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            
            logger.info("RAG 시스템 초기화 완료")
            
        except Exception as e:
            logger.error(f"RAG 시스템 초기화 실패: {e}")
            # RAG 없이도 작동하도록 설정
            self.retriever = None
    
    def build_knowledge_base(self):
        """기본 지식 베이스 구축"""
        try:
            # 기본 인구 관련 지식 문서들
            knowledge_docs = [
                {
                    "content": """
                    서울시 인구 현황 (2024년 기준)
                    - 총 인구: 약 950만명
                    - 인구 밀도: 15,000명/km²
                    - 주요 특징: 수도권 집중 현상, 고령화 진행
                    - 외국인 비율: 약 3.5%
                    """,
                    "metadata": {"source": "seoul_population", "type": "statistics"}
                },
                {
                    "content": """
                    강남구 지역 특성
                    - 총 인구: 약 55만명
                    - 면적: 39.5km²
                    - 인구 밀도: 14,000명/km²
                    - 주요 특징: IT 업계 중심지, 고소득층 집중, 외국인 비율 높음
                    - 주요 업종: IT, 금융, 교육, 의료
                    """,
                    "metadata": {"source": "gangnam_characteristics", "type": "regional_info"}
                },
                {
                    "content": """
                    한국 인구 정책 및 트렌드
                    - 저출산 고령화 대응 정책
                    - 외국인 유치 정책 (글로벌 인재)
                    - 수도권 집중 완화 정책
                    - 디지털 전환에 따른 인구 이동
                    - 코로나19 이후 원격근무 확산
                    """,
                    "metadata": {"source": "korea_policy", "type": "policy"}
                },
                {
                    "content": """
                    인구 분석 지표 해석 가이드
                    - 외국인 비율 10% 이상: 글로벌 기업 집중 지역
                    - 20-30대 비율 40% 이상: 젊은 층 집중 지역
                    - 인구 밀도 20,000명/km² 이상: 고밀도 도시 지역
                    - 주야간 인구 차이 30% 이상: 업무 중심 지역
                    """,
                    "metadata": {"source": "analysis_guide", "type": "interpretation"}
                },
                {
                    "content": """
                    강남구 행정동별 특성
                    - 역삼동: IT 기업 밀집, 외국인 많음
                    - 삼성동: 대기업 본사, 고급 주거지
                    - 청담동: 고급 상권, 주거지
                    - 신사동: 상권, 주거지 혼재
                    - 논현동: 상권 중심, 주거지
                    """,
                    "metadata": {"source": "gangnam_dongs", "type": "regional_detail"}
                }
            ]
            
            # 문서를 벡터 스토어에 추가
            texts = [doc["content"] for doc in knowledge_docs]
            metadatas = [doc["metadata"] for doc in knowledge_docs]
            
            self.vectorstore.add_texts(texts=texts, metadatas=metadatas)
            logger.info(f"기본 지식 베이스 구축 완료: {len(knowledge_docs)}개 문서")
            
        except Exception as e:
            logger.error(f"지식 베이스 구축 실패: {e}")
    
    def add_custom_knowledge(self, content: str, metadata: Dict = None):
        """사용자 정의 지식 추가"""
        try:
            if metadata is None:
                metadata = {"source": "custom", "type": "user_defined"}
            
            self.vectorstore.add_texts(texts=[content], metadatas=[metadata])
            logger.info("사용자 정의 지식 추가 완료")
            
        except Exception as e:
            logger.error(f"사용자 정의 지식 추가 실패: {e}")
    
    def retrieve_relevant_context(self, query: str, k: int = 3) -> List[str]:
        """관련 컨텍스트 검색"""
        try:
            if self.retriever is None:
                return []
            
            docs = self.retriever.get_relevant_documents(query)
            contexts = [doc.page_content for doc in docs[:k]]
            return contexts
            
        except Exception as e:
            logger.error(f"컨텍스트 검색 실패: {e}")
            return []

## ---------------- 문서 처리 시스템 ----------------------

def extract_text_from_file(file_path: str) -> str:
    """파일에서 텍스트 추출 (PDF, DOCX 지원)"""
    try:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            doc = fitz.open(file_path)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text
            
        elif ext == ".docx":
            doc = Document(file_path)
            text_parts = []
            
            # 단락 텍스트
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            
            # 표 데이터
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
            
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}. PDF 또는 DOCX만 허용됩니다.")
            
    except Exception as e:
        logger.error(f"파일 텍스트 추출 실패: {file_path}, 오류: {e}")
        raise

def extract_document_metadata(file_path: str) -> Dict[str, Any]:
    """문서 메타데이터 추출"""
    try:
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "created_date": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
            "modified_date": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
        }
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".docx":
            try:
                doc = Document(file_path)
                # 문서 속성 추출
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
                if doc.core_properties.created:
                    metadata["doc_created"] = doc.core_properties.created.isoformat()
                if doc.core_properties.modified:
                    metadata["doc_modified"] = doc.core_properties.modified.isoformat()
            except Exception as e:
                logger.warning(f"DOCX 메타데이터 추출 실패: {e}")
        
        return metadata
        
    except Exception as e:
        logger.error(f"메타데이터 추출 실패: {file_path}, 오류: {e}")
        return {}

def chunk_text(text: str, base_metadata: Dict) -> List[Dict]:
    """텍스트를 의미 있는 청크로 분할"""
    try:
        # 문단별 분할
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        chunks = []
        current_chunk = []
        current_length = 0
        max_chunk_length = 1000  # 문자 기준
        
        for i, para in enumerate(paragraphs):
            if current_length + len(para) > max_chunk_length and current_chunk:
                # 현재 청크 저장
                chunk_text = "\n".join(current_chunk)
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    "chunk_index": len(chunks),
                    "chunk_type": "paragraph_group",
                    "word_count": len(chunk_text.split()),
                    "source_type": "document"
                })
                
                chunks.append({
                    "text": chunk_text,
                    "metadata": chunk_metadata
                })
                
                # 새 청크 시작
                current_chunk = [para]
                current_length = len(para)
            else:
                current_chunk.append(para)
                current_length += len(para)
        
        # 마지막 청크 처리
        if current_chunk:
            chunk_text = "\n".join(current_chunk)
            chunk_metadata = base_metadata.copy()
            chunk_metadata.update({
                "chunk_index": len(chunks),
                "chunk_type": "paragraph_group",
                "word_count": len(chunk_text.split()),
                "source_type": "document"
            })
            
            chunks.append({
                "text": chunk_text,
                "metadata": chunk_metadata
            })
        
        return chunks
        
    except Exception as e:
        logger.error(f"텍스트 청킹 실패: {e}")
        return []

class DocumentRAGSystem(RAGSystem):
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.pdf', '.docx']
    
    def add_document(self, file_path: str, custom_metadata: Dict = None) -> bool:
        """문서 파일을 RAG 시스템에 추가"""
        try:
            # 텍스트 추출
            text_content = extract_text_from_file(file_path)
            
            # 메타데이터 추출
            doc_metadata = extract_document_metadata(file_path)
            
            # 사용자 정의 메타데이터 병합
            if custom_metadata:
                doc_metadata.update(custom_metadata)
            
            # 텍스트 청킹
            chunks = chunk_text(text_content, doc_metadata)
            
            if not chunks:
                logger.warning(f"문서에서 추출된 청크가 없습니다: {file_path}")
                return False
            
            # 벡터 스토어에 추가
            self.vectorstore.add_texts(
                texts=[chunk["text"] for chunk in chunks],
                metadatas=[chunk["metadata"] for chunk in chunks]
            )
            
            logger.info(f"문서 추가 완료: {file_path}, 청크 수: {len(chunks)}")
            return True
            
        except Exception as e:
            logger.error(f"문서 추가 실패: {file_path}, 오류: {e}")
            return False
    
    def get_document_list(self) -> List[Dict]:
        """등록된 문서 목록 조회"""
        try:
            if self.vectorstore is None:
                return []
            
            # 벡터 스토어에서 고유한 문서 정보 추출
            collection = self.vectorstore._collection
            results = collection.get()
            
            if not results or not results['metadatas']:
                return []
            
            # 문서별로 그룹화
            documents = {}
            for i, metadata in enumerate(results['metadatas']):
                if metadata and 'file_name' in metadata:
                    file_name = metadata['file_name']
                    if file_name not in documents:
                        documents[file_name] = {
                            'file_name': file_name,
                            'title': metadata.get('title', file_name),
                            'author': metadata.get('author', 'Unknown'),
                            'upload_date': metadata.get('upload_date', metadata.get('created_date', 'Unknown')),
                            'category': metadata.get('category', 'document'),
                            'chunk_count': 0
                        }
                    documents[file_name]['chunk_count'] += 1
            
            return list(documents.values())
            
        except Exception as e:
            logger.error(f"문서 목록 조회 실패: {e}")
            return []

# RAG 시스템 인스턴스 생성 (문서 지원 포함)
rag_system = DocumentRAGSystem()

## ---------------- 1단계 : 사전준비 ----------------------

# 1) State 선언 --------------------
class PopulationAnalysisState(TypedDict):
    # 고정 정보
    dong_name: str
    population_data: Dict[str, float]  # Double 타입 처리
    time_stats: List[Dict]
    gender_stats: Dict[str, float]     # Double 타입 처리
    age_stats: Dict[str, float]        # Double 타입 처리
    data_summary: str
    analysis_strategy: Dict[str, Dict]
    
    # RAG 관련 정보 추가
    rag_context: List[str]
    relevant_knowledge: str

    # 분석 로그
    current_analysis: str
    current_result: str
    current_step: str
    analysis_log: List[Dict[str, str]]
    evaluation: List[Dict[str, str]]
    next_step: str
    final_report: str

# 2) RAG 컨텍스트 검색 --------------------
def retrieve_rag_context(state: PopulationAnalysisState) -> PopulationAnalysisState:
    """분석에 필요한 RAG 컨텍스트 검색"""
    dong_name = state.get("dong_name", "")
    population_data = state.get("population_data", {})
    
    # 검색 쿼리 생성
    queries = [
        f"{dong_name} 인구 특성",
        f"강남구 {dong_name} 지역 정보",
        f"외국인 비율 {population_data.get('longForeigner', 0)}% 인구 분석",
        "서울시 생활인구 데이터 분석",
        "외국인과 내국인의 생활 패턴 분석"
    ]
    
    # RAG에서 관련 컨텍스트 검색
    all_contexts = []
    for query in queries:
        contexts = rag_system.retrieve_relevant_context(query, k=5)
        all_contexts.extend(contexts)
    
    # 중복 제거 및 정리
    unique_contexts = list(set(all_contexts))
    relevant_knowledge = "\n\n".join(unique_contexts[:8])  # 상위 8개만 사용
    
    state["rag_context"] = unique_contexts
    state["relevant_knowledge"] = relevant_knowledge
    
    logger.info(f"RAG 컨텍스트 검색 완료: {len(unique_contexts)}개 문서")
    return state

# 3) 데이터 분석 (RAG 강화) --------------------
def analyze_population_data(state: PopulationAnalysisState) -> PopulationAnalysisState:
    """인구 데이터를 분석하여 요약 생성 (RAG 컨텍스트 활용)"""
    dong_name = state.get("dong_name", "")
    population_data = state.get("population_data", {})
    time_stats = state.get("time_stats", [])
    gender_stats = state.get("gender_stats", {})
    age_stats = state.get("age_stats", {})
    relevant_knowledge = state.get("relevant_knowledge", "")
    
    if not population_data:
        logger.warning("population_data가 비어 있습니다. 기본값으로 진행합니다.")
        population_data = {"total": 0.0, "local": 0.0, "longForeigner": 0.0, "tempForeigner": 0.0}
    
    # 데이터 유효성 검사 (Double 타입 처리)
    total_pop = float(population_data.get('total', 0))
    local_pop = float(population_data.get('local', 0))
    long_foreigner = float(population_data.get('longForeigner', 0))
    temp_foreigner = float(population_data.get('tempForeigner', 0))
    
    logger.info(f"분석할 데이터 - 총인구: {int(total_pop):,}명, 내국인: {int(local_pop):,}명")
    
    # RAG 컨텍스트를 활용한 데이터 요약 생성
    llm = ChatOpenAI(
        model="gpt-4o-mini",
        temperature=0.1,
        max_tokens=500
    )
    
    summary_prompt = ChatPromptTemplate.from_template("""
    다음 인구 데이터와 관련 지식을 바탕으로 데이터 요약을 생성해주세요.
    
    [관련 지식]
    {relevant_knowledge}
    
    [인구 데이터]
    - 동 이름: {dong_name}
    - 총 인구: {total_pop:,}명
    - 내국인: {local_pop:,}명
    - 장기 외국인: {long_foreigner:,}명
    - 단기 외국인: {temp_foreigner:,}명
    - 외국인 비율: {foreigner_ratio:.1f}%
    
    [시간대별 데이터]
    {time_stats_summary}
    
    [성별 데이터]
    - 남성: {male_pop:,}명
    - 여성: {female_pop:,}명
    
    위 데이터를 바탕으로 다음을 분석해주세요:
    1. 전체적인 인구 규모와 구성
    2. 외국인 비율의 의미 (관련 지식 참고)
    3. 지역 특성과의 연관성
    4. 서울시/강남구 평균과의 비교
    
    간결하고 객관적으로 작성해주세요 (5줄 이내).
    """)
    
    # 데이터 계산
    foreigner_ratio = ((long_foreigner + temp_foreigner) / total_pop * 100) if total_pop > 0 else 0
    male_pop = float(gender_stats.get('male', 0))
    female_pop = float(gender_stats.get('female', 0))
    
    # 시간대별 데이터 요약
    time_stats_summary = ""
    if time_stats:
        peak_time = max(time_stats[:6], key=lambda x: float(x.get('totalPopulation', 0)))
        time_stats_summary = f"피크 시간대: {peak_time.get('timeRange', 'N/A')} ({peak_time.get('totalPopulation', 0):,}명)"
    
    # 요약 생성
    summary_chain = summary_prompt | llm | StrOutputParser()
    
    try:
        data_summary = summary_chain.invoke({
            "relevant_knowledge": relevant_knowledge,
            "dong_name": dong_name,
            "total_pop": int(total_pop),
            "local_pop": int(local_pop),
            "long_foreigner": int(long_foreigner),
            "temp_foreigner": int(temp_foreigner),
            "foreigner_ratio": foreigner_ratio,
            "time_stats_summary": time_stats_summary,
            "male_pop": int(male_pop),
            "female_pop": int(female_pop)
        })
        
        state["data_summary"] = data_summary
        logger.info("데이터 요약 생성 완료 (RAG 컨텍스트 활용)")
        
    except Exception as e:
        logger.error(f"데이터 요약 생성 실패: {e}")
        state["data_summary"] = f"{dong_name}의 총 인구는 {int(total_pop):,}명이며, 외국인 비율은 {foreigner_ratio:.1f}%입니다."
    
    return state

# 4) 분석 전략 수립 --------------------
def generate_analysis_strategy(state: PopulationAnalysisState) -> PopulationAnalysisState:
    """분석 전략 수립"""
    data_summary = state.get("data_summary", "")
    dong_name = state.get("dong_name", "")

    prompt = ChatPromptTemplate.from_template("""
당신은 인구 데이터 분석 전문가입니다. 주어진 {dong_name}의 '데이터 요약'을 기반으로, 가장 효과적인 분석 전략을 수립해야 합니다.
다음 세 가지 분석 부문별로 심층적인 분석 방향과 핵심 질문을 제시해 주세요.
모든 답변은 **제공된 데이터 요약 내용에 근거**하여 작성해야 합니다.

- 데이터 요약:
{data_summary}

아래 **명시된 딕셔너리 형식**을 정확히 준수하여 출력해 주세요. 다른 서식(예: JSON, 마크다운)은 사용하지 마세요.

{{
"인구구성분석": {{
"분석전략": "총 인구 대비 내외국인 비율, 남녀 성비, 연령대별 분포 등 인구 구성의 특징을 심층적으로 분석하여 지역의 핵심적인 인구 프로파일을 정의합니다.",
"핵심질문": [
"내국인과 외국인 생활인구의 시간대별 활동 반경 차이는 지역 상권의 업종 구성과 밀접하게 연관되어 있는가?",
"남녀 성비와 연령대별 인구 분포가 지역 내 편의시설 및 서비스 업종에 어떤 시사점을 주는가?"
]
}},
"시간패턴분석": {{
"분석전략": "요일별, 시간대별 인구 변동 패턴을 분석하여 특정 시간대에 인구가 급증하거나 감소하는 원인을 파악하고, 주요 활동 시간대를 정의합니다.",
"핵심질문": [
"특정 시간대에 유입되는 생활인구의 주요 활동 목적(쇼핑, 출퇴근, 여가 등)은 무엇이며, 이는 지역 편의시설의 운영 시간 및 서비스 구성에 어떤 시사점을 주는가?",
"특정 시간대의 인구 급증이 교통량, 대중교통 이용률에 어떤 영향을 미치는가?"
]
}},
"지역특성분석": {{
"분석전략": "주변 지역과의 인구 이동 패턴, 외부 유입 요인(관광, 상업시설 등)을 분석하여 해당 지역이 가진 고유한 특성(상업지, 주거지, 복합지 등)을 규명합니다.",
"핵심질문": [
"이 지역의 주요 인구 유입 및 유출 원인은 무엇이며, 이는 지역 경제 활성화에 어떤 영향을 미치는가?",
"인구 변화 패턴과 외부 유입 요인(예: 대형 쇼핑몰, 공원) 간의 상관관계는 어떻게 나타나는가?"
]
}}
}}
"""
)

    llm = ChatOpenAI(model="gpt-4o-mini")
    formatted_prompt = prompt.format(dong_name=dong_name, data_summary=data_summary)
    response = llm.invoke(formatted_prompt)

    # 딕셔너리로 변환
    dict_value = response.content.strip()
    if isinstance(dict_value, str):
        try:
            strategy_dict = ast.literal_eval(dict_value)
        except Exception as e:
            raise ValueError("analysis_strategy를 딕셔너리로 변환하는 데 실패했습니다.") from e

    return {
        **state,
        "analysis_strategy": strategy_dict
    }

# 5) 1단계 하나로 묶기 --------------------
def preProcessing_Analysis(dong_name: str, population_data: dict, time_stats: list, gender_stats: dict, age_stats: dict) -> PopulationAnalysisState:
    """전처리 단계"""
    # state 초기화
    initial_state: PopulationAnalysisState = {
        "dong_name": dong_name,
        "population_data": population_data,
        "time_stats": time_stats,
        "gender_stats": gender_stats,
        "age_stats": age_stats,
        "data_summary": '',
        "analysis_strategy": {},
        "rag_context": [],
        "relevant_knowledge": "",

        "current_analysis": '',
        "current_result": '',
        "current_step": '',
        "analysis_log": [],
        "evaluation": [],
        "next_step": '',
        "final_report": ''
    }

    # RAG 컨텍스트 검색
    initial_state = retrieve_rag_context(initial_state)

    # 데이터 분석
    state = analyze_population_data(initial_state)

    # 분석 전략 수립
    state = generate_analysis_strategy(state)

    # 첫번째 분석 시작
    analysis_strategy = state["analysis_strategy"]
    first_analysis = "인구구성분석"
    strategy_text = analysis_strategy[first_analysis]["분석전략"]
    core_questions = analysis_strategy[first_analysis]["핵심질문"]
    selected_question = random.choice(core_questions)

    return {
            **state,
            "current_analysis": selected_question,
            "current_step": first_analysis
            }

## ---------------- 2단계 : 분석 Agent ----------------------

# 1) 분석 실행 --------------------
def execute_analysis(state: PopulationAnalysisState) -> PopulationAnalysisState:
    """현재 분석 실행"""
    llm = ChatOpenAI(model="gpt-4o-mini")

    current_analysis = state.get("current_analysis", "")
    current_step = state.get("current_step", "")
    analysis_strategy = state.get("analysis_strategy", "")
    data_summary = state.get("data_summary", "")
    dong_name = state.get("dong_name", "")
    relevant_knowledge = state.get("relevant_knowledge", "")

    # 분석 전략 추출
    strategy_block = ""
    if isinstance(analysis_strategy, dict):
        strategy_block = analysis_strategy.get(current_step, {}).get("분석전략", "")
    elif isinstance(analysis_strategy, str):
        try:
            parsed = ast.literal_eval(analysis_strategy)
            strategy_block = parsed.get(current_step, {}).get("분석전략", "")
        except Exception:
            strategy_block = ""

    # RAG 컨텍스트를 활용한 강화된 프롬프트 구성
    prompt = ChatPromptTemplate.from_template("""
당신은 인구 데이터 분석 전문가입니다.
다음 정보와 관련 지식을 바탕으로 분석을 수행해주세요.

[관련 지식 베이스]
{relevant_knowledge}

[분석 대상 정보]
- 지역: {dong_name}
- 데이터 요약: {data_summary}
- 분석 전략({current_step}): {strategy}
- 분석 질문: {analysis}

위 정보를 바탕으로 분석 질문에 대해 구체적이고 실용적인 답변을 제공해주세요.
특히 관련 지식 베이스의 정보를 활용하여 다음을 고려해주세요:
1. 서울시/강남구 평균과의 비교
2. 지역 특성과의 연관성
3. 정책적 맥락과의 연결
4. 실용적인 인사이트 제공

답변은 3-4문장으로 간결하게 작성해주세요.
""")

    formatted_prompt = prompt.format(
        relevant_knowledge=relevant_knowledge,
        dong_name=dong_name,
        data_summary=data_summary,
        strategy=strategy_block,
        current_step=current_step,
        analysis=current_analysis
    )

    try:
        response = llm.invoke(formatted_prompt)
        analysis_result = response.content.strip()
        
        # (1) 분석 로그 저장 (질문/답변 1쌍)
        state["analysis_log"].append({"analysis": current_analysis, "result": analysis_result})
        
        logger.info(f"분석 완료 ({current_step}): {len(analysis_result)}자")
        
    except Exception as e:
        logger.error(f"분석 실행 실패: {e}")
        analysis_result = f"{dong_name}의 {current_step} 분석을 수행했습니다."
        state["analysis_log"].append({"analysis": current_analysis, "result": analysis_result})

    return {
        **state,
        "current_result": analysis_result
    }

# 2) 분석 평가 --------------------
def evaluate_analysis(state: PopulationAnalysisState) -> PopulationAnalysisState:
    """분석 결과 평가"""
    llm = ChatOpenAI(model="gpt-4o-mini")

    current_analysis = state.get("current_analysis", "")
    current_result = state.get("current_result", "")
    current_step = state.get("current_step", "")

    # 프롬프트 구성
    prompt = ChatPromptTemplate.from_template("""
당신은 인구 데이터 분석 평가자입니다.
다음 분석 결과를 평가해주세요.

[분석 정보]
- 분석 단계: {current_step}
- 분석 질문: {analysis}
- 분석 결과: {result}

위 정보를 바탕으로 아래 두 가지 항목에 따라 분석 결과를 평가해주세요.
- 분석의 구체성: 분석이 얼마나 구체적이고 실질적인 내용을 포함하고 있는지
- 실용성: 분석 결과가 실제 활용 가능한 인사이트를 제공하는지

각 항목에 대해 '상', '중', '하' 중 하나로 평가해주세요.

최종 결과는 아래 형식의 딕셔너리로만 출력해주세요:
{{
  "분석의 구체성": "상",
  "실용성": "중"
}}
""")

    formatted_prompt = prompt.format(
        current_step=current_step,
        analysis=current_analysis,
        result=current_result
    )

    response = llm.invoke(formatted_prompt)
    
    # 딕셔너리로 변환
    eval_dict = response.content.strip()
    if isinstance(eval_dict, str):
        try:
            eval_dict = ast.literal_eval(eval_dict)
        except Exception as e:
            eval_dict = {"분석의 구체성": "중", "실용성": "중"}

    # (2) 평가 저장 (인덱스 포함)
    evaluation = state.get("evaluation", [])
    eval_dict["analysis_index"] = len(state["analysis_log"]) - 1
    evaluation.append(eval_dict)

    return {
        **state,
        "evaluation": evaluation
    }

# 3) 다음 단계 결정 --------------------
def decide_next_step(state: PopulationAnalysisState) -> PopulationAnalysisState:
    """다음 분석 단계 결정"""
    evaluation = state.get("evaluation", [])
    analysis_log = state.get("analysis_log", [])
    current_step = state.get("current_step", "")

    # (1) 분석이 3회를 초과하면 종료
    if len(analysis_log) >= 3:
        next_step = "end"
    # (2) 현재 단계에 따라 다음 단계 결정
    elif current_step == "인구구성분석":
        next_step = "시간패턴분석"
    elif current_step == "시간패턴분석":
        next_step = "지역특성분석"
    else:
        next_step = "end"

    return {
        **state,
        "next_step": next_step
    }

# 4) 다음 분석 생성 --------------------
def generate_next_analysis(state: PopulationAnalysisState) -> PopulationAnalysisState:
    """다음 분석 질문 생성"""
    analysis_strategy = state.get("analysis_strategy", {})
    next_step = state.get("next_step", "")
    
    if next_step in analysis_strategy:
        core_questions = analysis_strategy[next_step]["핵심질문"]
        selected_question = random.choice(core_questions)
        
        return {
            **state,
            "current_analysis": selected_question,
            "current_step": next_step,
            "current_result": ""
        }
    
    return state

# 5) 최종 보고서 생성 --------------------
def generate_final_report(state: PopulationAnalysisState) -> PopulationAnalysisState:
    """최종 분석 보고서 생성 (RAG 컨텍스트 활용)"""
    dong_name = state.get("dong_name", "")
    analysis_log = state.get("analysis_log", [])
    relevant_knowledge = state.get("relevant_knowledge", "")
    
    # RAG 컨텍스트를 활용한 종합 분석
    llm = ChatOpenAI(
        model="gpt-4o-mini",
        temperature=0.1,
        max_tokens=800
    )
    
    # 분석 결과 요약
    analysis_summary = ""
    for i, log in enumerate(analysis_log):
        analysis_summary += f"{i+1}. {log['analysis']}\n   → {log['result']}\n\n"
    
    # RAG 컨텍스트를 활용한 종합 보고서 생성
    report_prompt = ChatPromptTemplate.from_template("""
다음 정보를 바탕으로 종합적인 인구 분석 보고서를 작성해주세요.

[관련 지식 베이스]
{relevant_knowledge}

[분석 대상]
- 지역: {dong_name}

[세부 분석 결과]
{analysis_summary}

위 정보를 종합하여 다음 형식으로 보고서를 작성해주세요:

🏙️ {dong_name} 인구 데이터 종합 분석 보고서

📊 주요 발견사항
(관련 지식을 바탕으로 한 핵심 인사이트 3-4개)

🔍 세부 분석
(각 분석 결과를 요약)

💡 실용적 제안
(데이터와 지식을 바탕으로 한 구체적 제안 2-3개)

📈 지역 특성 요약
(강남구 및 서울시 평균과의 비교)

⏰ 분석 일시: {timestamp}
""")
    
    try:
        report_chain = report_prompt | llm | StrOutputParser()
        final_report = report_chain.invoke({
            "relevant_knowledge": relevant_knowledge,
            "dong_name": dong_name,
            "analysis_summary": analysis_summary,
            "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
        
        logger.info("RAG 기반 최종 보고서 생성 완료")
        
    except Exception as e:
        logger.error(f"최종 보고서 생성 실패: {e}")
        # 기본 보고서 생성
        report_sections = []
        for i, log in enumerate(analysis_log):
            report_sections.append(f"""
{i+1}. {log['analysis']}
   → {log['result']}
""")
        
        final_report = f"""
🏙️ {dong_name} 인구 데이터 종합 분석 보고서

{''.join(report_sections)}

📊 분석 완료: {len(analysis_log)}개 항목
⏰ 분석 일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
    
    return {
        **state,
        "final_report": final_report
    }

# 6) Agent --------------------
# 분기 판단 함수
def route_next(state: PopulationAnalysisState) -> Literal["generate", "summarize"]:
    return "summarize" if state["next_step"] == "end" else "generate"

# 그래프 정의 시작
builder = StateGraph(PopulationAnalysisState)

# 노드 추가
builder.add_node("execute", execute_analysis)
builder.add_node("evaluate", evaluate_analysis)
builder.add_node("decide", decide_next_step)
builder.add_node("generate", generate_next_analysis)
builder.add_node("summarize", generate_final_report)

# 노드 연결
builder.set_entry_point("execute")
builder.add_edge("execute", "evaluate")
builder.add_edge("evaluate", "decide")
builder.add_conditional_edges("decide", route_next)
builder.add_edge("generate", "execute")      # 루프
builder.add_edge("summarize", END)           # 종료

# 컴파일
graph = builder.compile()

## ---------------- FastAPI 엔드포인트 ----------------------

@app.get("/")
async def root():
    return {"message": "인구 데이터 LangGraph Agent 분석 서버가 실행 중입니다", "status": "running"}

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "model": "gpt-4o-mini",
        "timestamp": datetime.now().isoformat()
    }

@app.post("/analyze/langgraph")
async def analyze_with_langgraph(data: Dict[str, Any]):
    """LangGraph Agent를 사용한 인구 데이터 분석 (RAG 통합)"""
    
    try:
        dong_name = data.get('dongName', '알 수 없는 동')
        population_data = data.get('populationData', {})
        time_stats = data.get('timeStats', [])
        gender_stats = data.get('genderStats', {})
        age_stats = data.get('ageStats', {})
        
        logger.info(f"{dong_name} LangGraph 분석 시작 (RAG 통합)...")
        logger.info(f"받은 인구 데이터: {population_data}")
        logger.info(f"받은 시간대 데이터 개수: {len(time_stats) if time_stats else 0}")
        logger.info(f"받은 성별 데이터: {gender_stats}")
        logger.info(f"받은 연령 데이터 타입: {type(age_stats)}, 내용: {age_stats}")
        
        # 전처리 단계 (RAG 컨텍스트 검색 포함)
        initial_state = preProcessing_Analysis(
            dong_name, population_data, time_stats, gender_stats, age_stats
        )
        
        # LangGraph 실행
        final_state = graph.invoke(initial_state)
        
        logger.info("LangGraph 분석 완료 (RAG 통합)")
        
        return {
            "status": "success",
            "dong_name": dong_name,
            "data_summary": final_state.get("data_summary", ""),
            "analysis_strategy": final_state.get("analysis_strategy", {}),
            "analysis_log": final_state.get("analysis_log", []),
            "evaluation": final_state.get("evaluation", []),
            "final_report": final_state.get("final_report", ""),
            "rag_context_count": len(final_state.get("rag_context", [])),
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"LangGraph 분석 실패: {e}")
        raise HTTPException(status_code=500, detail=f"분석 중 오류 발생: {str(e)}")

@app.post("/rag/add-knowledge")
async def add_custom_knowledge(data: Dict[str, Any]):
    """RAG 시스템에 사용자 정의 지식 추가"""
    try:
        content = data.get('content', '')
        metadata = data.get('metadata', {})
        
        if not content:
            raise HTTPException(status_code=400, detail="지식 내용이 필요합니다.")
        
        # RAG 시스템에 지식 추가
        rag_system.add_custom_knowledge(content, metadata)
        
        return {
            "status": "success",
            "message": "사용자 정의 지식이 성공적으로 추가되었습니다.",
            "content_length": len(content),
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"지식 추가 실패: {e}")
        raise HTTPException(status_code=500, detail=f"지식 추가 중 오류가 발생했습니다: {str(e)}")

@app.get("/rag/status")
async def get_rag_status():
    """RAG 시스템 상태 확인"""
    try:
        return {
            "status": "success",
            "rag_initialized": rag_system.retriever is not None,
            "vectorstore_ready": rag_system.vectorstore is not None,
            "embeddings_ready": rag_system.embeddings is not None,
            "supported_formats": rag_system.supported_formats,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"RAG 상태 확인 실패: {e}")
        return {
            "status": "error",
            "message": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.post("/rag/upload-document")
async def upload_document(
    file: UploadFile = File(...),
    title: str = Form(None),
    category: str = Form("document"),
    description: str = Form(None)
):
    """문서 파일 업로드 및 RAG 시스템에 추가"""
    try:
        # 파일 형식 검증
        if not file.filename.lower().endswith(('.pdf', '.docx')):
            raise HTTPException(
                status_code=400, 
                detail="지원하지 않는 파일 형식입니다. PDF 또는 DOCX만 허용됩니다."
            )
        
        # 임시 디렉토리 생성
        temp_dir = "./temp_documents"
        os.makedirs(temp_dir, exist_ok=True)
        
        # 임시 파일 저장
        temp_file_path = os.path.join(temp_dir, file.filename)
        
        with open(temp_file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # 사용자 정의 메타데이터
        custom_metadata = {
            "title": title or file.filename,
            "category": category,
            "description": description,
            "upload_date": datetime.now().isoformat()
        }
        
        # RAG 시스템에 추가
        success = rag_system.add_document(temp_file_path, custom_metadata)
        
        # 임시 파일 삭제
        try:
            os.remove(temp_file_path)
        except Exception as e:
            logger.warning(f"임시 파일 삭제 실패: {e}")
        
        if success:
            return {
                "status": "success",
                "message": "문서가 성공적으로 추가되었습니다.",
                "filename": file.filename,
                "metadata": custom_metadata,
                "timestamp": datetime.now().isoformat()
            }
        else:
            raise HTTPException(status_code=500, detail="문서 처리 중 오류가 발생했습니다.")
            
    except Exception as e:
        logger.error(f"문서 업로드 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/rag/documents")
async def get_documents():
    """등록된 문서 목록 조회"""
    try:
        documents = rag_system.get_document_list()
        return {
            "status": "success",
            "documents": documents,
            "count": len(documents),
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"문서 목록 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/rag/documents/{filename}")
async def delete_document(filename: str):
    """문서 삭제"""
    try:
        # 벡터 스토어에서 해당 문서의 청크들 삭제
        if rag_system.vectorstore is None:
            raise HTTPException(status_code=404, detail="RAG 시스템이 초기화되지 않았습니다.")
        
        collection = rag_system.vectorstore._collection
        results = collection.get()
        
        if not results or not results['metadatas']:
            raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")
        
        # 해당 파일의 청크들 찾기
        delete_ids = []
        for i, metadata in enumerate(results['metadatas']):
            if metadata and metadata.get('file_name') == filename:
                delete_ids.append(results['ids'][i])
        
        if delete_ids:
            collection.delete(ids=delete_ids)
            logger.info(f"문서 삭제 완료: {filename}, 삭제된 청크 수: {len(delete_ids)}")
            return {
                "status": "success",
                "message": f"문서 '{filename}'이 성공적으로 삭제되었습니다.",
                "deleted_chunks": len(delete_ids),
                "timestamp": datetime.now().isoformat()
            }
        else:
            raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")
            
    except Exception as e:
        logger.error(f"문서 삭제 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    
    logger.info("=" * 60)
    logger.info("🚀 인구 데이터 LangGraph Agent with RAG & Document Support 서버 시작")
    logger.info(f"📍 서버 주소: http://{HOST}:{PORT}")
    logger.info(f"🔧 환경: {ENVIRONMENT}")
    logger.info(f"🧠 모델: GPT-4o-mini")
    logger.info(f"📚 RAG: ChromaDB + Sentence Transformers")
    logger.info(f"📄 문서 지원: PDF, DOCX")
    logger.info("=" * 60)
    
    try:
        uvicorn.run(
            app,
            host=HOST,
            port=PORT,
            log_level="info",
            reload=False
        )
    except KeyboardInterrupt:
        logger.info("서버가 종료되었습니다.")
    except Exception as e:
        logger.error(f"서버 실행 중 오류 발생: {e}")

        logger.error(f"서버 실행 중 오류 발생: {e}")
